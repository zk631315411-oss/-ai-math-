from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\ai-math")
MATERIAL_DIR = ROOT / "星火杯参赛材料"
SOURCE = MATERIAL_DIR / "作品说明文档_智学助手_APA引用版.docx"
OUT = MATERIAL_DIR / "作品说明文档_智学助手_APA引用版_结构图优化版.docx"


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="9CB3C9", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def style_cell(cell, text, *, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, bold=False, color=None, fill="FFFFFF"):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    set_para(p, align=align, before=0, after=0, line=1.05)
    for run in p.runs:
        set_font(run, size=size, bold=bold, color=color)


def set_table_width(table, total_width):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total_width / 635)))

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    widths = [Cm(2.2), Cm(4.5), total_width - Cm(2.2) - Cm(4.5)]
    for col, width in zip(grid.gridCol_lst, widths):
        col.w = int(width / 635)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


doc = Document(SOURCE)
section = doc.sections[0]
usable_width = section.page_width - section.left_margin - section.right_margin

for para in doc.paragraphs:
    if para.text.startswith("应用效果预期。"):
        para.add_run(" 其结构如图1所示。")
        break

ref_anchor = next(p for p in doc.paragraphs if p.text.strip() == "参考文献")

caption = ref_anchor.insert_paragraph_before("图1 智学助手教育智能体结构")
set_para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=3, line=1.0)
for run in caption.runs:
    set_font(run, size=11, bold=True, color=(31, 78, 121))

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_width(table, usable_width)

headers = ["阶段", "模块", "作用说明"]
header_fill = "D9E8F4"
for i, text in enumerate(headers):
    style_cell(
        table.cell(0, i),
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10.5,
        bold=True,
        color=(31, 78, 121),
        fill=header_fill,
    )

rows = [
    ("1", "教材页锚定", "限定当前教材页、章节与知识边界，保证所有回答都围绕教材内容展开"),
    ("2", "学习者画像", "汇总掌握阶段、薄弱点和学习记录，形成可追踪的学习状态"),
    ("3", "认知阶段判断", "结合前置关系与最近发展区，判断学生当前适合接受的支持强度"),
    ("4", "导学答疑与练习", "提供分步解释、提示和按页练习，把讲解、练习和批改连成一体"),
    ("5", "反馈回写", "把新的提问、错因和掌握变化写回画像，更新后续学习建议"),
]

for data in rows:
    row = table.add_row().cells
    for idx, text in enumerate(data):
        align = WD_ALIGN_PARAGRAPH.CENTER if idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
        size = 10.2 if idx < 2 else 10.1
        style_cell(row[idx], text, align=align, size=size, fill="F7FBFF")

ref_anchor._p.addprevious(table._tbl)

note = doc.add_paragraph("注：图中各环节表示从教材约束到学习反馈的连续支持链条。")
set_para(note, align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=2, line=1.0)
for run in note.runs:
    set_font(run, size=9.2, color=(100, 100, 100))
table._tbl.addnext(note._p)

doc.save(OUT)
print(OUT)
