from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\ai-math")
MATERIAL_DIR = ROOT / "星火杯参赛材料"
SOURCE = MATERIAL_DIR / "作品说明文档_智学助手_APA引用版.docx"
OUT = MATERIAL_DIR / "作品说明文档_智学助手_APA引用版_末尾结构图版.docx"


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


doc = Document(SOURCE)
doc.add_page_break()

title = doc.add_paragraph()
set_paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
run = title.add_run("图1 智学助手教育智能体结构示意图")
set_font(run, size=14, bold=True, color=(31, 78, 121))

subtitle = doc.add_paragraph()
set_paragraph(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
run = subtitle.add_run("以教材约束为边界，以学习者状态为依据，以持续反馈促进理解深化")
set_font(run, size=10.2, color=(90, 90, 90))

page_width = doc.sections[0].page_width
left = doc.sections[0].left_margin
right = doc.sections[0].right_margin
usable = page_width - left - right
center = left + usable / 2

def add_shape(kind, x, y, w, h):
    shp = doc.inline_shapes if False else None


def add_rect(x, y, w, h, fill, line, text, title, body, title_color):
    shp = doc.Shapes.AddShape(1, x, y, w, h)
    shp.Fill.ForeColor.RGB = fill
    shp.Line.ForeColor.RGB = line
    shp.Line.Weight = 1.1
    shp.TextFrame.MarginLeft = 7
    shp.TextFrame.MarginRight = 7
    shp.TextFrame.MarginTop = 4
    shp.TextFrame.MarginBottom = 4
    shp.TextFrame.VerticalAnchor = 3
    shp.TextFrame.TextRange.Text = title + "\r" + body
    tr = shp.TextFrame.TextRange
    tr.ParagraphFormat.Alignment = 1
    tr.Font.NameFarEast = "仿宋"
    tr.Font.Name = "Times New Roman"
    tr.Font.Size = 9.1
    tr.Font.Color = (60, 60, 60)
    tr.Paragraphs(1).Range.Font.Size = 10.5
    tr.Paragraphs(1).Range.Font.Bold = True
    tr.Paragraphs(1).Range.Font.Color = title_color
    return shp


def add_band(x, y, w, h, text, fill, line, color):
    shp = doc.Shapes.AddShape(1, x, y, w, h)
    shp.Fill.ForeColor.RGB = fill
    shp.Line.ForeColor.RGB = line
    shp.Line.Weight = 1.2
    shp.TextFrame.MarginLeft = 8
    shp.TextFrame.MarginRight = 8
    shp.TextFrame.MarginTop = 3
    shp.TextFrame.MarginBottom = 3
    shp.TextFrame.VerticalAnchor = 3
    shp.TextFrame.TextRange.Text = text
    tr = shp.TextFrame.TextRange
    tr.ParagraphFormat.Alignment = 1
    tr.Font.NameFarEast = "仿宋"
    tr.Font.Name = "Times New Roman"
    tr.Font.Size = 10.5
    tr.Font.Bold = True
    tr.Font.Color = color
    return shp


def add_arrow(x, y, text="↓"):
    shp = doc.Shapes.AddTextbox(1, x, y, 24, 20)
    shp.Line.Visible = False
    shp.Fill.Visible = False
    tr = shp.TextFrame.TextRange
    tr.Text = text
    tr.ParagraphFormat.Alignment = 1
    tr.Font.NameFarEast = "仿宋"
    tr.Font.Name = "Times New Roman"
    tr.Font.Size = 18
    tr.Font.Bold = True
    tr.Font.Color = 3940090
    return shp


blue_fill = 15656936
green_fill = 15925611
orange_fill = 16771782
header_fill = 1983737
grey_fill = 15658741

blue_line = 9160037
green_line = 10252349
orange_line = 14278301
grey_line = 12041634

blue_text = 1983737
green_text = 4424722
orange_text = 9200925
grey_text = 3947593

add_band(center - 215, 100, 430, 28, "教育目标：精准诊断   适切支架   持续反馈", header_fill, header_fill, 16777215)
add_arrow(center - 12, 133)

box_w = 400
box_x = center - box_w / 2
top = 160
step = 79

add_rect(box_x, top, box_w, 48, blue_fill, blue_line, "", "一　学习入口\n教材阅读\n文字提问\n截图求助", "一　学习入口", blue_text)
add_arrow(center - 12, top + 53)
add_rect(box_x, top + step, box_w, 48, green_fill, green_line, "", "二　知识基础\n页码锚定\n知识图谱\n前置关系", "二　知识基础", green_text)
add_arrow(center - 12, top + step + 53)
add_rect(box_x, top + step * 2, box_w, 48, orange_fill, orange_line, "", "三　学习者诊断\n掌握阶段\n薄弱概念\n学习记录", "三　学习者诊断", orange_text)
add_arrow(center - 12, top + step * 2 + 53)
add_rect(box_x, top + step * 3, box_w, 48, blue_fill, blue_line, "", "四　导学支持\n分步讲解\n提示追问\n支架调节", "四　导学支持", blue_text)
add_arrow(center - 12, top + step * 3 + 53)
add_rect(box_x, top + step * 4, box_w, 48, green_fill, green_line, "", "五　练习评价\n按页出题\n过程批改\n错因分析", "五　练习评价", green_text)
add_arrow(center - 12, top + step * 4 + 53)
add_rect(box_x, top + step * 5, box_w, 48, orange_fill, orange_line, "", "六　反馈更新\n画像回写\n补弱建议\n策略调整", "六　反馈更新", orange_text)

add_band(center - 215, top + step * 5 + 72, 430, 28, "评价证据：提问质量   概念变化   错因分布   后续建议", grey_fill, grey_line, grey_text)

note = doc.add_paragraph()
set_paragraph(note, align=WD_ALIGN_PARAGRAPH.CENTER, before=6)
run = note.add_run("注：直线箭头表示学习支持的推进方向，反馈更新用于形成下一轮学习建议。")
set_font(run, size=9.1, color=(95, 95, 95))

doc.save(OUT)
print(OUT)
