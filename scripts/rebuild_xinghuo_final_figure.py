from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\ai-math")
MATERIAL_DIR = ROOT / "星火杯参赛材料"
SOURCE = MATERIAL_DIR / "作品说明文档_智学助手_APA引用版.docx"
OUT = MATERIAL_DIR / "作品说明文档_智学助手_APA引用版_结构图重做版.docx"


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.05):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_margins(cell, top=120, start=130, bottom=120, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, value="nil", color="FFFFFF", size="0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), value)
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def clear_cell(cell, fill="FFFFFF"):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, fill)
    set_cell_border(cell, "nil")
    set_cell_margins(cell, 40, 40, 40, 40)
    set_paragraph(cell.paragraphs[0], align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.0)


def style_box(cell, title, body, fill, border, title_color):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, fill)
    set_cell_border(cell, "single", border, "12")
    set_cell_margins(cell, 150, 160, 150, 160)

    p = cell.paragraphs[0]
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.05)
    title_run = p.add_run(title)
    set_run_font(title_run, size=10.8, bold=True, color=title_color)
    p.add_run("\n")
    body_run = p.add_run(body)
    set_run_font(body_run, size=9.2, color=(50, 50, 50))


def style_band(cell, text, fill, border, color=(255, 255, 255)):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, fill)
    set_cell_border(cell, "single", border, "12")
    set_cell_margins(cell, 150, 160, 150, 160)
    p = cell.paragraphs[0]
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.05)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=color)


def style_arrow(cell, text):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "FFFFFF")
    set_cell_border(cell, "nil")
    set_cell_margins(cell, 40, 40, 40, 40)
    p = cell.paragraphs[0]
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=18, bold=True, color=(59, 94, 126))


def set_table_widths(table):
    widths = [Inches(1.95), Inches(0.32), Inches(1.95), Inches(0.32), Inches(1.95)]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w / 635) for w in widths)))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.w = int(width / 635)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


doc = Document(SOURCE)

# Keep the existing report intact and place the redesigned figure at the very end.
doc.add_page_break()

title = doc.add_paragraph("附图  智学助手教育智能体结构示意图")
set_paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line=1.0)
for run in title.runs:
    set_run_font(run, size=14, bold=True, color=(31, 78, 121))

subtitle = doc.add_paragraph("以教材约束为边界，以学习者状态为依据，以持续反馈促进理解深化")
set_paragraph(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line=1.0)
for run in subtitle.runs:
    set_run_font(run, size=10.2, color=(85, 85, 85))

table = doc.add_table(rows=7, cols=5)
set_table_widths(table)

for row in table.rows:
    for cell in row.cells:
        clear_cell(cell)

goal = table.cell(0, 0).merge(table.cell(0, 4))
style_band(goal, "教育目标：精准诊断　适切支架　持续反馈", "1F4E79", "1F4E79")

style_arrow(table.cell(1, 2), "↓")

style_box(
    table.cell(2, 0),
    "一　学习入口",
    "教材阅读\n文字提问\n截图求助",
    "EAF3F8",
    "8BAFC5",
    (31, 78, 121),
)
style_arrow(table.cell(2, 1), "→")
style_box(
    table.cell(2, 2),
    "二　知识基础",
    "页码锚定\n知识图谱\n前置关系",
    "F3F8EA",
    "9BB77D",
    (67, 112, 50),
)
style_arrow(table.cell(2, 3), "→")
style_box(
    table.cell(2, 4),
    "三　学习者诊断",
    "掌握阶段\n薄弱概念\n学习记录",
    "FFF4E6",
    "D8AA64",
    (139, 89, 29),
)

style_arrow(table.cell(3, 4), "↓")

style_box(
    table.cell(4, 4),
    "四　导学支持",
    "分步讲解\n提示追问\n支架调节",
    "EAF3F8",
    "8BAFC5",
    (31, 78, 121),
)
style_arrow(table.cell(4, 3), "←")
style_box(
    table.cell(4, 2),
    "五　练习评价",
    "按页出题\n过程批改\n错因分析",
    "F3F8EA",
    "9BB77D",
    (67, 112, 50),
)
style_arrow(table.cell(4, 1), "←")
style_box(
    table.cell(4, 0),
    "六　反馈更新",
    "画像回写\n补弱建议\n策略调整",
    "FFF4E6",
    "D8AA64",
    (139, 89, 29),
)

style_arrow(table.cell(5, 0), "↓")

evidence = table.cell(6, 0).merge(table.cell(6, 4))
style_band(
    evidence,
    "评价证据：提问质量　概念变化　错因分布　后续建议",
    "EEF2F5",
    "B8C7D3",
    color=(60, 60, 60),
)

note = doc.add_paragraph("注：图中直线箭头表示学习支持的主要推进方向，反馈更新用于形成下一轮学习建议。")
set_paragraph(note, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=0, line=1.0)
for run in note.runs:
    set_run_font(run, size=9.2, color=(95, 95, 95))

doc.save(OUT)
print(OUT)
